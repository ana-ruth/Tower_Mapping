from docx import Document
from .mapFetcher import *    
#from mapFetcher import *

document = Document()


def generate_report(df):
    
    df.apply(tower_summary, axis = 1)
    document.save("sample_report.docx")


def tower_summary(row):
    
    document.add_paragraph("Tower Name: "+ str(row["tower_name"]))

    document.add_paragraph("Address: "+str(row["address"]))

    document.add_paragraph("Latitude: "+str(row["latitude"]))

    document.add_paragraph("Longitude: "+ str(row["longitude"]))

    map_img = map_generator(row["latitude"], row["longitude"])
    
    document.add_picture(map_img)

